from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.core.exceptions import PermissionDenied

from django.http import HttpResponse
from docx import Document
import os
from .decoradores import supervisor_required

# Definir la ruta base para guardar los archivos
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS_DIR = os.path.join(BASE_DIR, 'documentos')

# Asegurarse de que el directorio de documentos existe
if not os.path.exists(DOCS_DIR):
    os.makedirs(DOCS_DIR)

@login_required
def index(request):
    return render(request, 'gestionusuario/index.html')

@login_required
def profile(request):
    return render(request, 'gestionusuario/profile.html')

@login_required
def grupo(request):
    return render(request, 'gestionusuario/grupo.html')

@supervisor_required
def crearart(request):
    if request.method == 'POST':
        supervisor = request.POST.get('supervisor')
        empresa = request.POST.get('empresa')
        gerencia = request.POST.get('gerencia')
        fecha = request.POST.get('fecha')
        hora_inicio = request.POST.get('hora_inicio')
        hora_termino = request.POST.get('hora_termino')
        superintendencia = request.POST.get('superintendencia')
        lugar_especifico = request.POST.get('lugar_especifico')
        trabajo_a_realizar = request.POST.get('trabajo_a_realizar')

        # Crear documento Word
        doc = Document()
        doc.add_heading('Datos del Formulario', 0)
        doc.add_paragraph(f'Supervisor: {supervisor}')
        doc.add_paragraph(f'Empresa: {empresa}')
        doc.add_paragraph(f'Gerencia: {gerencia}')
        doc.add_paragraph(f'Fecha: {fecha}')
        doc.add_paragraph(f'Hora inicio: {hora_inicio}')
        doc.add_paragraph(f'Hora término: {hora_termino}')
        doc.add_paragraph(f'Superintendencia/Dirección: {superintendencia}')
        doc.add_paragraph(f'Lugar específico: {lugar_especifico}')
        doc.add_paragraph(f'Trabajo a realizar: {trabajo_a_realizar}')

        doc_path = os.path.join(DOCS_DIR, 'datos_formulario_1.docx')
        doc.save(doc_path)

        return redirect('crearart2')

    return render(request, 'gestionusuario/crearart.html')

@supervisor_required
def crearart2(request):
    if request.method == 'POST':
        pregunta_1 = request.POST.get('pregunta_1')
        pregunta_2 = request.POST.get('pregunta_2')
        pregunta_3 = request.POST.get('pregunta_3')
        pregunta_4 = request.POST.get('pregunta_4')
        pregunta_5 = request.POST.get('pregunta_5')
        pregunta_6 = request.POST.get('pregunta_6')

        doc_path = os.path.join(DOCS_DIR, 'datos_formulario_1.docx')
        if os.path.exists(doc_path):
            doc = Document(doc_path)
        else:
            return HttpResponse("El archivo datos_formulario_1.docx no se encontró.")

        doc.add_heading('Preguntas Adicionales', level=1)
        doc.add_paragraph(f'Pregunta 1: {pregunta_1}')
        doc.add_paragraph(f'Pregunta 2: {pregunta_2}')
        doc.add_paragraph(f'Pregunta 3: {pregunta_3}')
        doc.add_paragraph(f'Pregunta 4: {pregunta_4}')
        doc.add_paragraph(f'Pregunta 5: {pregunta_5}')
        doc.add_paragraph(f'Pregunta 6: {pregunta_6}')

        doc_path = os.path.join(DOCS_DIR, 'datos_formulario_2.docx')
        doc.save(doc_path)

        return redirect('crearart3')

    return render(request, 'gestionusuario/crearart2.html')

@supervisor_required
def crearart3(request):
    if request.method == 'POST':
        nombre = request.POST.get('nombre', '')

        doc_path = os.path.join(DOCS_DIR, 'datos_formulario_1.docx')
        if os.path.exists(doc_path):
            doc = Document(doc_path)
        else:
            return HttpResponse("El archivo datos_formulario_1.docx no se encontró.")

        for i in range(1, 31):
            opcion_seleccionada = request.POST.get(f'opcion_{i}', '')

            print(f"Pregunta {i}:")
            print(f"  - Opción seleccionada: {opcion_seleccionada}")

            doc.add_paragraph(f'aguante nisha nixyeva : {i}')

        doc_path = os.path.join(DOCS_DIR, 'datos_formulario_3.docx')
        try:
            doc.save(doc_path)
        except Exception as e:
            return HttpResponse(f"Error al guardar el documento: {str(e)}")

        return redirect('crearart4')

    return render(request, 'gestionusuario/crearart3.html')

@supervisor_required
def crearart4(request):
    if request.method == 'POST':
        riesgos = []
        medidas_control = []

        doc_path = os.path.join(DOCS_DIR, 'datos_formulario_3.docx')
        if os.path.exists(doc_path):
            doc = Document(doc_path)
        else:
            return HttpResponse("El archivo datos_formulario_3.docx no se encontró.")

        doc.add_heading('Riesgos y Medidas de Control', level=1)

        for i in range(1, 9):
            riesgo = request.POST.get(f'riesgo{i}')
            medida_control = request.POST.get(f'medida_control{i}')
            riesgos.append(riesgo)
            medidas_control.append(medida_control)
            doc.add_paragraph(f'Riesgo {i}: {riesgo}')
            doc.add_paragraph(f'Medida de Control {i}: {medida_control}')

        doc_path = os.path.join(DOCS_DIR, 'datos_formulario_4.docx')
        try:
            doc.save(doc_path)
        except Exception as e:
            return HttpResponse(f"Error al guardar el documento: {str(e)}")

        return redirect('crearart5')

    context = {'range': range(8)}
    return render(request, 'gestionusuario/crearart4.html', context)

@supervisor_required
def crearart5(request):
    if request.method == 'POST':
        pregunta_1 = request.POST.get('pregunta_1')
        pregunta_2 = request.POST.get('pregunta_2')
        pregunta_3 = request.POST.get('pregunta_3')
        pregunta_4 = request.POST.get('pregunta_4')
        pregunta_5 = request.POST.get('pregunta_5')
        pregunta_6 = request.POST.get('pregunta_6')
        pregunta_7 = request.POST.get('pregunta_7')
        pregunta_8 = request.POST.get('pregunta_8')

        doc_path = os.path.join(DOCS_DIR, 'datos_formulario_4.docx')
        if os.path.exists(doc_path):
            doc = Document(doc_path)
        else:
            return HttpResponse("El archivo datos_formulario_4.docx no se encontró.")

        doc.add_heading('Preguntas Finales', level=1)
        doc.add_paragraph(f'Pregunta 1: {pregunta_1}')
        doc.add_paragraph(f'Pregunta 2: {pregunta_2}')
        doc.add_paragraph(f'Pregunta 3: {pregunta_3}')
        doc.add_paragraph(f'Pregunta 4: {pregunta_4}')
        doc.add_paragraph(f'Pregunta 5: {pregunta_5}')
        doc.add_paragraph(f'Pregunta 6: {pregunta_6}')
        doc.add_paragraph(f'Pregunta 7: {pregunta_7}')
        doc.add_paragraph(f'Pregunta 8: {pregunta_8}')

        doc_path = os.path.join(DOCS_DIR, 'datos_formulario_final.docx')
        try:
            doc.save(doc_path)
        except Exception as e:
            return HttpResponse(f"Error al guardar el documento: {str(e)}")

        return redirect('alguna_vista')

    return render(request, 'gestionusuario/crearart5.html')
